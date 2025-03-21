from flask import Flask, request, jsonify, send_from_directory
from flask_cors import CORS
from werkzeug.utils import secure_filename
from PIL import Image
from PyPDF2 import PdfMerger
from reportlab.pdfgen import canvas
import pandas as pd
import numpy as np
import os
import uuid
import time
from paddleocr import PaddleOCR
from io import BytesIO
from pdf2docx import Converter
import difflib
import zipfile
import py7zr
import shutil
import PyPDF2
import threading

app = Flask(__name__)
CORS(app)

# 全局配置临时文件夹，确保它存在
app.config['TEMP_FOLDER'] = os.path.join(os.path.dirname(__file__), 'temp')
os.makedirs(app.config['TEMP_FOLDER'], exist_ok=True)

# 配置文件上传
UPLOAD_FOLDER = 'uploads'
OUTPUT_FOLDER = 'outputs'
ALLOWED_EXTENSIONS = {'jpg', 'jpeg', 'png', 'gif', 'pdf', 'docx', 'xlsx', 'txt', 'zip', '7z'}
MAX_CONTENT_LENGTH = 20 * 1024 * 1024  # 20MB

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['OUTPUT_FOLDER'] = OUTPUT_FOLDER

# 确保目录存在
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

# 存储上传的图片信息
uploads = {}

# 存储上传的文件信息
compression_files = {}

# 初始化OCR模型
ocr = PaddleOCR(use_angle_cls=True, lang="ch", use_gpu=False)

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

# =============== 图片转Excel功能 ===============

@app.route('/image-to-excel', methods=['POST'])
def image_to_excel():
    if 'image' not in request.files:
        return jsonify({'success': False, 'message': '未找到图片'}), 400
    
    file = request.files['image']
    if file.filename == '':
        return jsonify({'success': False, 'message': '未选择文件'}), 400
    
    if file and allowed_file(file.filename):
        # 生成唯一文件名
        filename = secure_filename(file.filename)
        timestamp = int(time.time())
        unique_id = f"{timestamp}_{uuid.uuid4().hex[:8]}"
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], f"{unique_id}_{filename}")
        file.save(file_path)
        
        try:
            # 使用PaddleOCR识别图片中的表格
            result = ocr.ocr(file_path, cls=True)
            
            # 提取文本和位置信息
            boxes = []
            texts = []
            for line in result:
                for item in line:
                    box = item[0]  # 坐标
                    text = item[1][0]  # 文本内容
                    boxes.append(box)
                    texts.append(text)
            
            # 根据位置信息排列成表格
            # 提取y坐标作为行标识
            y_centers = [sum([box[1][1], box[2][1]]) / 2 for box in boxes]
            
            # 聚类y坐标，将接近的坐标归为一行
            y_tolerance = 20  # 容差值，可根据实际情况调整
            rows = []
            current_row = [0]
            for i in range(1, len(y_centers)):
                if abs(y_centers[i] - y_centers[current_row[0]]) < y_tolerance:
                    current_row.append(i)
                else:
                    rows.append(current_row)
                    current_row = [i]
            if current_row:
                rows.append(current_row)
            
            # 对每行内的元素按x坐标排序
            table_data = []
            for row_indices in rows:
                # 提取该行的所有文本和对应的x坐标
                row_boxes = [boxes[i] for i in row_indices]
                row_texts = [texts[i] for i in row_indices]
                
                # 计算x中心坐标
                x_centers = [sum([box[0][0], box[1][0]]) / 2 for box in row_boxes]
                
                # 根据x坐标排序
                sorted_indices = np.argsort(x_centers)
                sorted_texts = [row_texts[i] for i in sorted_indices]
                
                table_data.append(sorted_texts)
            
            # 创建Excel文件
            df = pd.DataFrame(table_data)
            
            # 确保所有列有相同数量的元素
            max_cols = max([len(row) for row in table_data]) if table_data else 0
            
            # 生成Excel文件
            excel_name = f"excel_{unique_id}.xlsx"
            excel_path = os.path.join(app.config['OUTPUT_FOLDER'], excel_name)
            df.to_excel(excel_path, index=False, header=False)
            
            # 生成下载URL
            excel_url = f"https://excel.sube.top/download/{excel_name}"
            
            # 返回结果
            return jsonify({
                'success': True,
                'data': {
                    'rows': len(table_data),
                    'columns': max_cols,
                    'excelUrl': excel_url
                }
            })
            
        except Exception as e:
            return jsonify({'success': False, 'message': f'处理表格失败: {str(e)}'}), 500
    
    return jsonify({'success': False, 'message': '不支持的文件类型'}), 400

# =============== 图片转PDF功能 ===============

@app.route('/image-to-pdf-upload', methods=['POST'])
def upload_image_for_pdf():
    try:
        if 'image' not in request.files:
            response = jsonify({'success': False, 'message': '没有上传图片'})
            response.headers['Content-Type'] = 'application/json'
            return response, 400
        
        image_file = request.files['image']
        if not image_file or not image_file.filename:
            response = jsonify({'success': False, 'message': '无效的图片文件'})
            response.headers['Content-Type'] = 'application/json'
            return response, 400
        
        # 安全处理文件名
        original_filename = image_file.filename
        filename = secure_filename(original_filename)
        
        # 创建唯一的文件名避免冲突
        timestamp = int(time.time())
        unique_filename = f"{timestamp}_{filename}"
        
        # 保存图片文件
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], unique_filename)
        image_file.save(filepath)
        
        # 返回成功响应
        result = {
            'success': True,
            'data': {
                'originalFilename': original_filename,
                'filename': unique_filename,
                'url': f"https://excel.sube.top/uploads/{unique_filename}"
            }
        }
        
        response = jsonify(result)
        response.headers['Content-Type'] = 'application/json'
        return response
    
    except Exception as e:
        error_message = f"上传图片失败: {str(e)}"
        print(error_message)
        response = jsonify({'success': False, 'message': error_message})
        response.headers['Content-Type'] = 'application/json'
        return response, 500

@app.route('/merge-pdf', methods=['POST'])
def merge_pdf():
    try:
        data = request.json
        image_ids = data.get('imageIds', [])
        
        if not image_ids or not isinstance(image_ids, list) or len(image_ids) == 0:
            response = jsonify({'success': False, 'message': '图片ID或文件名无效'})
            response.headers['Content-Type'] = 'application/json'
            return response, 400
        
        # 尝试两种方式查找图片：
        # 1. 通过uploads字典中的ID（旧方法）
        # 2. 通过文件名直接查找（新方法）
        images = []
        
        # 首先尝试通过uploads字典查找
        for img_id in image_ids:
            if img_id in uploads:
                images.append({'id': img_id, 'path': uploads[img_id]['path'], 'index': uploads[img_id]['index']})
                
        # 如果没有找到图片，尝试通过文件名查找
        if not images:
            print("通过uploads字典未找到图片，尝试通过文件名查找")
            for filename in image_ids:
                # 检查是否是文件名
                if isinstance(filename, str) and ('.' in filename):
                    # 构建完整路径
                    file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                    # 检查文件是否存在
                    if os.path.exists(file_path):
                        # 根据字符串在列表中的位置确定索引
                        index = image_ids.index(filename)
                        images.append({'id': f"file-{index}", 'path': file_path, 'index': index})
        
        if not images:
            response = jsonify({'success': False, 'message': '未找到有效图片'})
            response.headers['Content-Type'] = 'application/json'
            return response, 400
        
        # 按索引排序图片
        images.sort(key=lambda x: x['index'])
        
        # 生成PDF
        pdf_name = f"pdf-{int(time.time())}.pdf"
        pdf_path = os.path.join(app.config['OUTPUT_FOLDER'], pdf_name)
        
        # 使用PdfMerger合并所有图片生成的PDF
        merger = PdfMerger()
        
        for img_info in images:
            # 将图片转换为PDF
            img_path = img_info['path']
            temp_pdf = BytesIO()
            
            try:
                # 打开图片
                image = Image.open(img_path)
                
                # 确定PDF页面大小
                width, height = image.size
                
                # 创建一个PDF页面
                c = canvas.Canvas(temp_pdf, pagesize=(width, height))
                c.drawImage(img_path, 0, 0, width, height)
                c.save()
                
                # 将BytesIO重置到开始位置
                temp_pdf.seek(0)
                
                # 将这个页面添加到合并器
                merger.append(temp_pdf)
            except Exception as e:
                print(f"处理图片出错: {str(e)}")
                # 继续处理下一张图片
                continue
        
        # 写入合并后的PDF
        merger.write(pdf_path)
        merger.close()
        
        # 计算文件大小
        file_size = os.path.getsize(pdf_path)
        file_size_kb = round(file_size / 1024, 2)
        
        # 构建PDF下载URL
        pdf_url = f"https://excel.sube.top/download/{pdf_name}"
        
        # 返回结果
        result = {
            'success': True,
            'data': {
                'pdfUrl': pdf_url,
                'fileSize': f"{file_size_kb} KB"
            }
        }
        
        response = jsonify(result)
        response.headers['Content-Type'] = 'application/json'
        return response
        
    except Exception as e:
        error_msg = f"合并PDF失败: {str(e)}"
        print(error_msg)
        response = jsonify({'success': False, 'message': error_msg})
        response.headers['Content-Type'] = 'application/json'
        return response, 500

# =============== PDF转Word功能 ===============

@app.route('/pdf-to-word', methods=['POST'])
def pdf_to_word():
    if 'pdf' not in request.files:
        return jsonify({'success': False, 'message': '未找到PDF文件'}), 400
    
    file = request.files['pdf']
    if file.filename == '':
        return jsonify({'success': False, 'message': '未选择文件'}), 400
    
    # 检查文件类型
    if not file.filename.lower().endswith('.pdf'):
        return jsonify({'success': False, 'message': '请上传PDF格式的文件'}), 400
    
    # 生成唯一文件名
    filename = secure_filename(file.filename)
    timestamp = int(time.time())
    unique_id = f"{timestamp}_{uuid.uuid4().hex[:8]}"
    pdf_path = os.path.join(app.config['UPLOAD_FOLDER'], f"{unique_id}_{filename}")
    file.save(pdf_path)
    
    try:
        # 生成Word文件名
        word_name = f"word_{unique_id}.docx"
        word_path = os.path.join(app.config['OUTPUT_FOLDER'], word_name)
        
        # 转换PDF为Word
        cv = Converter(pdf_path)
        cv.convert(word_path, start=0, end=None)
        cv.close()
        
        # 获取文件大小
        file_size = os.path.getsize(word_path)
        file_size_kb = round(file_size / 1024, 2)
        
        # 生成下载URL
        word_url = f"https://excel.sube.top/download/{word_name}"
        
        # 返回结果
        return jsonify({
            'success': True,
            'data': {
                'wordUrl': word_url,
                'fileSize': f"{file_size_kb}KB",
                'fileName': word_name
            }
        })
        
    except Exception as e:
        return jsonify({'success': False, 'message': f'转换失败: {str(e)}'}), 500

# =============== MD转Word功能 ===============

@app.route('/md-to-word', methods=['POST'])
def md_to_word():
    if 'md' not in request.files:
        return jsonify({'success': False, 'message': '未找到Markdown文件'}), 400
    
    file = request.files['md']
    if file.filename == '':
        return jsonify({'success': False, 'message': '未选择文件'}), 400
    
    # 检查文件类型
    if not file.filename.lower().endswith(('.md', '.markdown')):
        return jsonify({'success': False, 'message': '请上传Markdown格式的文件'}), 400
    
    # 生成唯一文件名
    filename = secure_filename(file.filename)
    timestamp = int(time.time())
    unique_id = f"{timestamp}_{uuid.uuid4().hex[:8]}"
    md_path = os.path.join(app.config['UPLOAD_FOLDER'], f"{unique_id}_{filename}")
    file.save(md_path)
    
    try:
        # 生成Word文件名
        word_name = f"word_{unique_id}.docx"
        word_path = os.path.join(app.config['OUTPUT_FOLDER'], word_name)
        
        # 使用python-docx将Markdown转换为Word
        from markdown import markdown
        from docx import Document
        from docx.shared import Pt, Inches
        from bs4 import BeautifulSoup

        # 读取Markdown内容
        with open(md_path, 'r', encoding='utf-8') as f:
            md_content = f.read()
        
        # 将Markdown转换为HTML
        html = markdown(md_content, extensions=['tables', 'fenced_code', 'codehilite'])
        
        # 使用BeautifulSoup解析HTML
        soup = BeautifulSoup(html, 'html.parser')
        
        # 创建Word文档
        doc = Document()
        
        # 设置文档页面边距
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # 处理HTML元素
        for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'ul', 'ol', 'pre', 'table']):
            if element.name.startswith('h'):
                level = int(element.name[1])
                heading = doc.add_heading(element.get_text(), level=level)
            elif element.name == 'p':
                p = doc.add_paragraph(element.get_text())
            elif element.name == 'ul':
                for li in element.find_all('li'):
                    p = doc.add_paragraph(li.get_text(), style='List Bullet')
            elif element.name == 'ol':
                for li in element.find_all('li'):
                    p = doc.add_paragraph(li.get_text(), style='List Number')
            elif element.name == 'pre':
                code = element.get_text()
                p = doc.add_paragraph(code)
                for run in p.runs:
                    run.font.name = 'Courier New'
                    run.font.size = Pt(10)
            elif element.name == 'table':
                rows = element.find_all('tr')
                if rows:
                    # 获取表格列数
                    cols = max([len(row.find_all(['td', 'th'])) for row in rows])
                    
                    # 创建表格
                    table = doc.add_table(rows=len(rows), cols=cols)
                    table.style = 'Table Grid'
                    
                    # 填充表格
                    for i, row in enumerate(rows):
                        cells = row.find_all(['td', 'th'])
                        for j, cell in enumerate(cells):
                            if j < cols:  # 确保不越界
                                table.cell(i, j).text = cell.get_text().strip()
        
        # 保存文档
        doc.save(word_path)
        
        # 获取文件大小
        file_size = os.path.getsize(word_path)
        file_size_kb = round(file_size / 1024, 2)
        
        # 生成下载URL
        word_url = f"https://excel.sube.top/download/{word_name}"
        
        # 返回结果
        return jsonify({
            'success': True,
            'data': {
                'wordUrl': word_url,
                'fileSize': f"{file_size_kb}KB",
                'fileName': word_name
            }
        })
        
    except Exception as e:
        return jsonify({'success': False, 'message': f'转换失败: {str(e)}'}), 500

# =============== 文件下载功能 ===============

@app.route('/download/<filename>', methods=['GET'])
def download_file(filename):
    try:
        # 检查文件是否存在
        file_path = os.path.join(app.config['OUTPUT_FOLDER'], filename)
        if not os.path.exists(file_path):
            response = jsonify({
                'success': False,
                'message': '文件不存在'
            })
            response.headers['Content-Type'] = 'application/json'
            return response, 404
        
        # 根据文件类型设置Content-Type
        mime_types = {
            '.pdf': 'application/pdf',
            '.xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            '.zip': 'application/zip',
            '.7z': 'application/x-7z-compressed',
            '.html': 'text/html',
            '.txt': 'text/plain',
            '.jpg': 'image/jpeg',
            '.jpeg': 'image/jpeg',
            '.png': 'image/png',
            '.gif': 'image/gif'
        }
        
        # 获取文件扩展名
        file_ext = os.path.splitext(filename)[1].lower()
        content_type = mime_types.get(file_ext, 'application/octet-stream')
        
        # 发送文件并设置正确的Content-Type
        response = send_from_directory(
            app.config['OUTPUT_FOLDER'], 
            filename, 
            as_attachment=True
        )
        response.headers['Content-Type'] = content_type
        response.headers['Content-Disposition'] = f'attachment; filename="{filename}"'
        response.headers['Access-Control-Allow-Origin'] = '*'
        return response
        
    except Exception as e:
        error_msg = f"下载文件失败: {str(e)}"
        print(error_msg)
        response = jsonify({
            'success': False,
            'message': error_msg
        })
        response.headers['Content-Type'] = 'application/json'
        return response, 500

# =============== 文本比较功能 ===============

@app.route('/compare-text', methods=['POST'])
def compare_text():
    try:
        # 获取请求中的文本
        data = request.json
        if not data or 'text1' not in data or 'text2' not in data:
            response = jsonify({'success': False, 'message': '请提供两段要比较的文本'})
            response.headers['Content-Type'] = 'application/json'
            return response, 400
        
        text1 = data['text1']
        text2 = data['text2']
        
        # 使用difflib的SequenceMatcher进行字符级别的比较
        matcher = difflib.SequenceMatcher(None, text1, text2)
        
        # 为并排显示准备数据
        text1_formatted = []
        text2_formatted = []
        
        # 优化处理连续的相同内容，合并为一个块
        for tag, i1, i2, j1, j2 in matcher.get_opcodes():
            # 处理相同部分的显示
            if tag == 'equal':
                # 如果内容少于2个字符并且不是空白符，则不单独显示
                if i2 - i1 <= 2 and not text1[i1:i2].isspace():
                    # 检查前后是否有内容
                    if text1_formatted and text1_formatted[-1]['class'] == 'unchanged':
                        text1_formatted[-1]['text'] += text1[i1:i2]
                    else:
                        text1_formatted.append({'text': text1[i1:i2], 'class': 'unchanged'})
                    
                    if text2_formatted and text2_formatted[-1]['class'] == 'unchanged':
                        text2_formatted[-1]['text'] += text2[j1:j2]
                    else:
                        text2_formatted.append({'text': text2[j1:j2], 'class': 'unchanged'})
                else:
                    text1_formatted.append({'text': text1[i1:i2], 'class': 'unchanged'})
                    text2_formatted.append({'text': text2[j1:j2], 'class': 'unchanged'})
            # 处理删除部分
            elif tag == 'delete':
                text1_formatted.append({'text': text1[i1:i2], 'class': 'removed'})
            # 处理添加部分
            elif tag == 'insert':
                text2_formatted.append({'text': text2[j1:j2], 'class': 'added'})
            # 处理替换部分
            elif tag == 'replace':
                text1_formatted.append({'text': text1[i1:i2], 'class': 'removed'})
                text2_formatted.append({'text': text2[j1:j2], 'class': 'added'})
        
        # 计算统计数据
        similarity = round(matcher.ratio() * 100, 2)
        added_chars = sum(j2 - j1 for tag, i1, i2, j1, j2 in matcher.get_opcodes() if tag in ('insert', 'replace'))
        removed_chars = sum(i2 - i1 for tag, i1, i2, j1, j2 in matcher.get_opcodes() if tag in ('delete', 'replace'))
        unchanged_chars = sum(i2 - i1 for tag, i1, i2, j1, j2 in matcher.get_opcodes() if tag == 'equal')
        total_chars = len(text1) + len(text2)
        different_chars = added_chars + removed_chars
        
        # 直接返回结果，包含差异内容
        result = {
            'success': True,
            'data': {
                'similarity': similarity,
                'added_chars': added_chars,
                'removed_chars': removed_chars,
                'unchanged_chars': unchanged_chars,
                'different_chars': different_chars,
                'total_chars': total_chars,
                'text1_formatted': text1_formatted,
                'text2_formatted': text2_formatted
            }
        }
        
        response = jsonify(result)
        response.headers['Content-Type'] = 'application/json'
        return response
        
    except Exception as e:
        error_msg = f"比较失败: {str(e)}"
        print(error_msg)
        response = jsonify({'success': False, 'message': error_msg})
        response.headers['Content-Type'] = 'application/json'
        return response, 500

# =============== 文件压缩功能 ===============

@app.route('/upload-file', methods=['POST'])
def upload_file():
    try:
        if 'file' not in request.files:
            return jsonify({'success': False, 'message': '未找到文件'}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({'success': False, 'message': '未选择文件'}), 400
        
        # 获取文件索引和原始文件名
        index = int(request.form.get('index', 0))
        original_filename = request.form.get('fileName', file.filename)
        
        # 生成唯一ID
        file_id = f"{int(time.time())}-{uuid.uuid4().hex[:8]}"
        
        # 安全处理文件名
        filename = secure_filename(file.filename)
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], f"{file_id}_{filename}")
        
        # 保存文件
        file.save(file_path)
        
        # 存储文件信息
        compression_files[file_id] = {
            'path': file_path,
            'originalName': original_filename,
            'index': index
        }
        
        return jsonify({
            'success': True,
            'data': {
                'fileId': file_id,
                'fileName': original_filename,
                'size': os.path.getsize(file_path)
            }
        })
    
    except Exception as e:
        print(f"上传文件错误: {str(e)}")
        return jsonify({'success': False, 'message': '服务器错误'}), 500

@app.route('/compress', methods=['POST'])
def compress_files():
    try:
        data = request.json
        format_type = data.get('format', 'zip')
        
        if not compression_files:
            return jsonify({'success': False, 'message': '未找到有效文件'}), 400
        
        # 创建临时目录存放文件
        temp_dir = os.path.join(app.config['UPLOAD_FOLDER'], f"temp_{int(time.time())}")
        os.makedirs(temp_dir, exist_ok=True)
        
        # 复制文件到临时目录，保持原来的文件名
        file_list = []
        for file_id, file_info in compression_files.items():
            src_path = file_info['path']
            dst_name = file_info['originalName']
            dst_path = os.path.join(temp_dir, dst_name)
            
            # 如果有重名，则添加序号
            if os.path.exists(dst_path):
                name_parts = os.path.splitext(dst_name)
                dst_name = f"{name_parts[0]}_{file_id[-4:]}{name_parts[1]}"
                dst_path = os.path.join(temp_dir, dst_name)
            
            shutil.copy2(src_path, dst_path)
            file_list.append({'name': dst_name, 'path': dst_path})
        
        # 生成压缩文件名
        timestamp = int(time.time())
        archive_name = f"files_{timestamp}"
        
        if format_type == 'zip':
            # 创建ZIP文件
            archive_path = os.path.join(app.config['OUTPUT_FOLDER'], f"{archive_name}.zip")
            with zipfile.ZipFile(archive_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
                for file_item in file_list:
                    zipf.write(file_item['path'], arcname=file_item['name'])
            
            file_extension = 'zip'
        else:
            # 创建7Z文件
            archive_path = os.path.join(app.config['OUTPUT_FOLDER'], f"{archive_name}.7z")
            with py7zr.SevenZipFile(archive_path, 'w') as szf:
                for file_item in file_list:
                    szf.write(file_item['path'], arcname=file_item['name'])
            
            file_extension = '7z'
        
        # 清理临时目录
        shutil.rmtree(temp_dir, ignore_errors=True)
        
        # 清空文件列表
        compression_files.clear()
        
        # 计算文件大小
        file_size = os.path.getsize(archive_path)
        file_size_kb = round(file_size / 1024, 2)
        
        if file_size_kb < 1024:
            file_size_str = f"{file_size_kb} KB"
        else:
            file_size_str = f"{round(file_size_kb / 1024, 2)} MB"
        
        # 构建下载URL
        file_url = f"https://excel.sube.top/download/{archive_name}.{file_extension}"
        
        return jsonify({
            'success': True,
            'data': {
                'fileSize': file_size_str,
                'fileName': f"{archive_name}.{file_extension}",
                'fileUrl': file_url
            }
        })
    
    except Exception as e:
        print(f"压缩文件错误: {str(e)}")
        return jsonify({'success': False, 'message': f'压缩失败: {str(e)}'}), 500

# =============== 根据URL合并PDF功能 ===============

@app.route('/urls-to-pdf', methods=['POST'])
def urls_to_pdf():
    try:
        # 获取请求数据
        data = request.json
        image_urls = data.get('imageUrls', [])
        
        if not image_urls or not isinstance(image_urls, list) or len(image_urls) == 0:
            response = jsonify({'success': False, 'message': '图片URL无效'})
            response.headers['Content-Type'] = 'application/json'
            return response, 400
        
        # 准备图片路径列表
        image_paths = []
        for url in image_urls:
            # 从URL中提取文件名
            filename = url.split('/')[-1]
            if not filename:
                continue
                
            # 构建本地路径
            file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            
            # 检查文件是否存在
            if os.path.exists(file_path):
                image_paths.append({'path': file_path, 'filename': filename})
        
        if not image_paths:
            response = jsonify({'success': False, 'message': '未找到有效图片文件'})
            response.headers['Content-Type'] = 'application/json'
            return response, 400
        
        # 生成PDF
        pdf_name = f"pdf-{int(time.time())}.pdf"
        pdf_path = os.path.join(app.config['OUTPUT_FOLDER'], pdf_name)
        
        # 使用PdfMerger合并所有图片生成的PDF
        merger = PdfMerger()
        
        for img_info in image_paths:
            # 将图片转换为PDF
            img_path = img_info['path']
            temp_pdf = BytesIO()
            
            try:
                # 打开图片
                image = Image.open(img_path)
                
                # 确定PDF页面大小
                width, height = image.size
                
                # 创建一个PDF页面
                c = canvas.Canvas(temp_pdf, pagesize=(width, height))
                c.drawImage(img_path, 0, 0, width, height)
                c.save()
                
                # 将BytesIO重置到开始位置
                temp_pdf.seek(0)
                
                # 将这个页面添加到合并器
                merger.append(temp_pdf)
            except Exception as e:
                print(f"处理图片出错: {str(e)}")
                # 继续处理下一张图片
                continue
        
        # 写入合并后的PDF
        merger.write(pdf_path)
        merger.close()
        
        # 计算文件大小
        file_size = os.path.getsize(pdf_path)
        file_size_kb = round(file_size / 1024, 2)
        
        # 构建PDF下载URL
        pdf_url = f"https://excel.sube.top/download/{pdf_name}"
        
        # 返回结果
        result = {
            'success': True,
            'data': {
                'pdfUrl': pdf_url,
                'fileSize': f"{file_size_kb} KB"
            }
        }
        
        response = jsonify(result)
        response.headers['Content-Type'] = 'application/json'
        return response
        
    except Exception as e:
        error_msg = f"根据URL合并PDF失败: {str(e)}"
        print(error_msg)
        response = jsonify({'success': False, 'message': error_msg})
        response.headers['Content-Type'] = 'application/json'
        return response, 500

# =============== PDF拆分合并功能 ===============

@app.route('/api/pdf/upload', methods=['POST'])
def pdf_upload():
    try:
        # 检查是否收到文件
        if 'file' not in request.files:
            return jsonify({
                'success': False,
                'message': '未接收到文件'
            }), 400
        
        file = request.files['file']
        
        # 检查文件是否为PDF
        if not file.filename.lower().endswith('.pdf'):
            return jsonify({
                'success': False,
                'message': '只支持PDF文件'
            }), 400
        
        # 生成唯一文件ID
        file_id = f"{int(time.time())}-{uuid.uuid4().hex[:8]}"
        
        # 保存文件到临时目录
        temp_path = os.path.join(app.config['TEMP_FOLDER'], f"{file_id}.pdf")
        file.save(temp_path)
        
        # 处理上传后的操作
        operation = request.form.get('operation', '')
        index = request.form.get('index', '0')
        
        return jsonify({
            'success': True,
            'fileId': file_id,
            'operation': operation,
            'index': index
        })
        
    except Exception as e:
        error_msg = f"PDF上传失败: {str(e)}"
        print(error_msg)
        return jsonify({
            'success': False,
            'message': error_msg
        }), 500

@app.route('/api/pdf/analyze', methods=['POST'])
def pdf_analyze():
    try:
        # 检查是否有文件
        if 'file' not in request.files:
            return jsonify({
                "success": False,
                "message": "未上传文件"
            }), 400
            
        file = request.files['file']
        
        # 检查文件类型
        if not file.filename.lower().endswith('.pdf'):
            return jsonify({
                "success": False,
                "message": "请上传PDF文件"
            }), 400
            
        # 生成文件ID
        file_id = f"{int(time.time())}-{uuid.uuid4().hex[:8]}"
        
        # 保存文件
        filename = secure_filename(file.filename)
        file_path = os.path.join(app.config['TEMP_FOLDER'], f"{file_id}_{filename}")
        file.save(file_path)
        
        # 分析PDF页数
        try:
            # 兼容PyPDF2不同版本
            try:
                # 新版本的PyPDF2
                with open(file_path, 'rb') as pdf_file:
                    pdf = PyPDF2.PdfReader(pdf_file)
                    page_count = len(pdf.pages)
            except:
                # 旧版本的PyPDF2
                with open(file_path, 'rb') as pdf_file:
                    pdf = PyPDF2.PdfFileReader(pdf_file)
                    page_count = pdf.getNumPages()
                    
            return jsonify({
                "success": True,
                "fileId": file_id,
                "pageCount": page_count
            }), 200
            
        except Exception as e:
            return jsonify({
                "success": False,
                "message": f"无法分析PDF: {str(e)}"
            }), 400
            
    except Exception as e:
        print(f"PDF分析错误: {str(e)}")
        return jsonify({
            "success": False,
            "message": f"处理PDF时发生错误: {str(e)}"
        }), 500

@app.route('/api/pdf/merge', methods=['POST'])
def pdf_merge():
    try:
        # 获取请求数据
        data = request.json
        if not data or not data.get('files'):
            return jsonify({
                'success': False,
                'message': '未提供文件ID列表'
            }), 400
        
        file_ids = data.get('files', [])
        output_filename = data.get('filename', 'merged.pdf')
        
        if not file_ids:
            return jsonify({
                'success': False,
                'message': '文件列表为空'
            }), 400
            
        # 合并PDF
        merger = PyPDF2.PdfMerger()
        
        for file_id in file_ids:
            pdf_path = os.path.join(app.config['TEMP_FOLDER'], f"{file_id}.pdf")
            if os.path.exists(pdf_path):
                merger.append(pdf_path)
            else:
                return jsonify({
                    'success': False,
                    'message': f'文件不存在: {file_id}.pdf'
                }), 404
        
        # 生成输出文件名
        output_id = str(uuid.uuid4())
        output_path = os.path.join(app.config['OUTPUT_FOLDER'], f"{output_id}.pdf")
        
        # 写入合并后的PDF
        merger.write(output_path)
        merger.close()
        
        # 获取文件大小
        file_size = os.path.getsize(output_path)
        file_size_str = format_file_size(file_size)
        
        # 生成下载URL
        pdf_url = f"{request.host_url.rstrip('/')}/download/{output_id}.pdf"
        
        return jsonify({
            'success': True,
            'pdfUrl': pdf_url,
            'fileSize': file_size_str
        })
        
    except Exception as e:
        error_msg = f"PDF合并失败: {str(e)}"
        print(error_msg)
        return jsonify({
            'success': False,
            'message': error_msg
        }), 500

@app.route('/api/pdf/split', methods=['POST'])
def pdf_split():
    try:
        # 获取文件ID和范围
        data = request.json
        if not data or 'fileId' not in data:
            return jsonify({"error": "缺少文件ID"}), 400
            
        file_id = data['fileId']
        ranges = data.get('ranges', [])
        
        if not ranges or len(ranges) == 0:
            return jsonify({"error": "缺少拆分范围"}), 400
            
        # 查找文件
        file_path = None
        for root, dirs, files in os.walk(app.config['TEMP_FOLDER']):
            for file in files:
                if file.startswith(file_id):
                    file_path = os.path.join(root, file)
                    break
            if file_path:
                break
                
        if not file_path:
            return jsonify({"error": "文件不存在或已过期"}), 404
            
        # 创建输出目录
        output_dir = os.path.join(app.config['TEMP_FOLDER'], f"split_{int(time.time())}")
        os.makedirs(output_dir, exist_ok=True)
        
        # 打开PDF文件
        try:
            # 兼容PyPDF2不同版本
            try:
                # 新版本的PyPDF2
                with open(file_path, 'rb') as pdf_file:
                    pdf = PyPDF2.PdfReader(pdf_file)
                    total_pages = len(pdf.pages)
                    
                    # 处理每个范围
                    output_files = []
                    for i, page_range in enumerate(ranges):
                        start = page_range.get('start', 0)
                        end = page_range.get('end', 0)
                        
                        if start < 1 or end > total_pages or start > end:
                            return jsonify({"error": f"范围 {i+1} ({start}-{end}) 无效，PDF共有 {total_pages} 页"}), 400
                        
                        output = PyPDF2.PdfWriter()
                        
                        # 添加页面
                        for page_num in range(start-1, end):
                            output.add_page(pdf.pages[page_num])
                            
                        # 保存拆分后的PDF
                        output_filename = f"split_{i+1}_{start}-{end}.pdf"
                        output_path = os.path.join(output_dir, output_filename)
                        
                        with open(output_path, 'wb') as output_file:
                            output.write(output_file)
                            
                        output_files.append(output_path)
            except:
                # 旧版本的PyPDF2
                with open(file_path, 'rb') as pdf_file:
                    pdf = PyPDF2.PdfFileReader(pdf_file)
                    total_pages = pdf.getNumPages()
                    
                    # 处理每个范围
                    output_files = []
                    for i, page_range in enumerate(ranges):
                        start = page_range.get('start', 0)
                        end = page_range.get('end', 0)
                        
                        if start < 1 or end > total_pages or start > end:
                            return jsonify({"error": f"范围 {i+1} ({start}-{end}) 无效，PDF共有 {total_pages} 页"}), 400
                        
                        output = PyPDF2.PdfFileWriter()
                        
                        # 添加页面
                        for page_num in range(start-1, end):
                            output.addPage(pdf.getPage(page_num))
                            
                        # 保存拆分后的PDF
                        output_filename = f"split_{i+1}_{start}-{end}.pdf"
                        output_path = os.path.join(output_dir, output_filename)
                        
                        with open(output_path, 'wb') as output_file:
                            output.write(output_file)
                            
                        output_files.append(output_path)
            
            # 创建ZIP文件
            timestamp = int(time.time())
            zip_filename = f"split_pdf_{timestamp}.zip"
            zip_path = os.path.join(app.config['OUTPUT_FOLDER'], zip_filename)
            
            with zipfile.ZipFile(zip_path, 'w') as zipf:
                for file in output_files:
                    zipf.write(file, os.path.basename(file))
            
            # 返回下载链接
            download_url = f"/download/{zip_filename}"
            return jsonify({
                "success": True,
                "zipUrl": request.host_url.rstrip('/') + download_url,
                "fileCount": len(output_files),
                "totalSize": format_file_size(os.path.getsize(zip_path))
            }), 200
        
        except Exception as e:
            return jsonify({
                "success": False,
                "message": f"处理PDF时出错: {str(e)}"
            }), 400
        
    except Exception as e:
        print(f"PDF拆分错误: {str(e)}")
        return jsonify({
            "success": False,
            "message": f"处理PDF时发生错误: {str(e)}"
        }), 500

# 格式化文件大小的辅助函数
def format_file_size(size_bytes):
    """将字节数转换为可读的文件大小字符串"""
    if size_bytes < 1024:
        return f"{size_bytes} B"
    elif size_bytes < 1024 * 1024:
        return f"{size_bytes/1024:.2f} KB"
    elif size_bytes < 1024 * 1024 * 1024:
        return f"{size_bytes/(1024*1024):.2f} MB"
    else:
        return f"{size_bytes/(1024*1024*1024):.2f} GB"

# 原before_first_request函数改为普通函数
def ensure_temp_folder_exists():
    # 确保临时文件夹存在
    if not os.path.exists(app.config['TEMP_FOLDER']):
        os.makedirs(app.config['TEMP_FOLDER'])
    if not os.path.exists(app.config['OUTPUT_FOLDER']):
        os.makedirs(app.config['OUTPUT_FOLDER'])
    if not os.path.exists(app.config['UPLOAD_FOLDER']):
        os.makedirs(app.config['UPLOAD_FOLDER'])

# 添加清理文件的函数
def cleanup_files():
    """定期清理超过5分钟的上传和输出文件"""
    while True:
        try:
            print("开始清理临时文件...")
            current_time = time.time()
            timeout = 5 * 60  # 5分钟 (300秒)
            
            # 清理uploads目录
            cleanup_directory(app.config['UPLOAD_FOLDER'], current_time, timeout)
            
            # 清理outputs目录
            cleanup_directory(app.config['OUTPUT_FOLDER'], current_time, timeout)
            
            # 清理temp目录
            cleanup_directory(app.config['TEMP_FOLDER'], current_time, timeout)
            
            print(f"文件清理完成，将在{timeout}秒后再次清理")
        except Exception as e:
            print(f"清理文件时出错: {str(e)}")
        
        # 等待5分钟再次执行
        time.sleep(timeout)

def cleanup_directory(directory, current_time, timeout):
    """清理指定目录中超时的文件"""
    if not os.path.exists(directory):
        return
    
    count = 0
    for filename in os.listdir(directory):
        file_path = os.path.join(directory, filename)
        
        # 跳过目录
        if os.path.isdir(file_path):
            # 只清理temp_开头的临时目录
            if filename.startswith('temp_'):
                try:
                    # 获取目录修改时间
                    modified_time = os.path.getmtime(file_path)
                    if current_time - modified_time > timeout:
                        shutil.rmtree(file_path, ignore_errors=True)
                        count += 1
                        print(f"已删除临时目录: {filename}")
                except Exception as e:
                    print(f"删除目录 {filename} 出错: {str(e)}")
            continue
        
        try:
            # 获取文件修改时间
            modified_time = os.path.getmtime(file_path)
            
            # 如果文件修改时间超过timeout秒，则删除
            if current_time - modified_time > timeout:
                os.remove(file_path)
                count += 1
        except Exception as e:
            print(f"处理文件 {filename} 时出错: {str(e)}")
    
    if count > 0:
        print(f"从 {directory} 删除了 {count} 个过期文件")

# 启动清理线程
def start_cleanup_thread():
    cleanup_thread = threading.Thread(target=cleanup_files, daemon=True)
    cleanup_thread.start()
    print("文件自动清理线程已启动，文件生命周期为5分钟")

# 在应用启动前直接调用
ensure_temp_folder_exists()
start_cleanup_thread()

# =============== 启动服务 ===============

if __name__ == '__main__':
    # 确保所有必要的文件夹都存在
    with app.app_context():
        ensure_temp_folder_exists()
    app.run(host='0.0.0.0', port=5000, debug=False) 